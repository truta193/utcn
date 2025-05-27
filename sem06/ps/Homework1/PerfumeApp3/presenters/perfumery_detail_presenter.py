import os
from sqlalchemy.orm import Session
from typing import List, Dict, Any, Optional
from datetime import datetime
import csv
from docx import Document

from models.repository.perfumery_repository import PerfumeryRepository
from models.repository.inventory_repository import InventoryRepository
from presenters.interfaces import PerfumeryDetailViewInterface
from schemas.perfumery_schemas import PerfumeryCreate, PerfumeryUpdate
from config import EXPORT_STORAGE_PATH


class PerfumeryDetailPresenter:

    def __init__(self, db: Session, view: PerfumeryDetailViewInterface):
        self.db = db
        self.view = view
        self.perfumery_repository = PerfumeryRepository(db)
        self.inventory_repository = InventoryRepository(db)
        self.export_dir = EXPORT_STORAGE_PATH
        os.makedirs(self.export_dir, exist_ok=True)

    def show_perfumery(self, perfumery_id: int):
        try:
            perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            inventory = self.inventory_repository.get_all_by_perfumery(perfumery_id)

            perfumery_data = self._convert_perfumery_to_dict(perfumery)
            inventory_data = [self._convert_inventory_to_dict(inv) for inv in inventory]

            return self.view.display_perfumery(perfumery_data, inventory_data)

        except Exception as e:
            self.view.display_error(str(e))
            return self.view.redirect_to_list()

    def filter_perfumery_inventory(self, perfumery_id: int, filter_type: str, filter_value: str):
        try:
            perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            if filter_type == 'availability':
                is_available = filter_value.lower() == 'true'
                inventory = self.inventory_repository.get_by_availability(perfumery_id, is_available)
                filter_description = "Available Items" if is_available else "Out of Stock Items"
            else:
                inventory = self.inventory_repository.get_all_by_perfumery(perfumery_id)
                filter_description = None

            perfumery_data = self._convert_perfumery_to_dict(perfumery)
            inventory_data = [self._convert_inventory_to_dict(inv) for inv in inventory]

            return self.view.display_filtered_inventory(perfumery_data, inventory_data, filter_description)

        except Exception as e:
            self.view.display_error(str(e))
            return self.view.redirect_to_list()

    def show_create_form(self):
        return self.view.show_create_form()

    def create_perfumery(self, form_data: Dict[str, Any]):
        try:
            perfumery_data = PerfumeryCreate(
                name=form_data['name'],
                address=form_data['address'],
                city=form_data['city'],
                phone=form_data.get('phone'),
                email=form_data.get('email'),
                manager_name=form_data.get('manager_name')
            )

            new_perfumery = self.perfumery_repository.create(
                name=perfumery_data.name,
                address=perfumery_data.address,
                city=perfumery_data.city,
                phone=perfumery_data.phone,
                email=perfumery_data.email,
                manager_name=perfumery_data.manager_name
            )

            self.view.display_success("Perfumery created successfully!")
            return self.view.redirect_to_detail(new_perfumery.id)

        except Exception as e:
            self.view.display_error(f"Error creating perfumery: {str(e)}")
            return self.view.show_create_form()

    def show_edit_form(self, perfumery_id: int):
        try:
            perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            perfumery_data = self._convert_perfumery_to_dict(perfumery)

            return self.view.show_edit_form(perfumery_data)

        except Exception as e:
            self.view.display_error(str(e))
            return self.view.redirect_to_list()

    def update_perfumery(self, perfumery_id: int, form_data: Dict[str, Any]):
        try:
            db_perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not db_perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            update_data = {}
            if form_data.get('name'):
                update_data['name'] = form_data['name']
            if form_data.get('address'):
                update_data['address'] = form_data['address']
            if form_data.get('city'):
                update_data['city'] = form_data['city']
            if 'phone' in form_data:
                update_data['phone'] = form_data['phone']
            if 'email' in form_data:
                update_data['email'] = form_data['email']
            if 'manager_name' in form_data:
                update_data['manager_name'] = form_data['manager_name']

            perfumery_update = PerfumeryUpdate(**update_data)
            updated_perfumery = self.perfumery_repository.update(perfumery_id,
                                                                 **perfumery_update.dict(exclude_unset=True))

            if not updated_perfumery:
                self.view.display_error("Failed to update perfumery")
                return self.view.redirect_to_list()

            self.view.display_success("Perfumery updated successfully!")
            return self.view.redirect_to_detail(perfumery_id)

        except Exception as e:
            self.view.display_error(f"Error updating perfumery: {str(e)}")
            return self.view.show_edit_form(self._convert_perfumery_to_dict(db_perfumery))

    def delete_perfumery(self, perfumery_id: int):
        try:
            db_perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not db_perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            success = self.perfumery_repository.delete(perfumery_id)

            if not success:
                self.view.display_error("Failed to delete perfumery")
            else:
                self.view.display_success("Perfumery deleted successfully!")

            return self.view.redirect_to_list()

        except Exception as e:
            self.view.display_error(f"Error deleting perfumery: {str(e)}")
            return self.view.redirect_to_list()

    def export_out_of_stock(self, perfumery_id: int, format_type: str):
        try:
            perfumery = self.perfumery_repository.get_by_id(perfumery_id)
            if not perfumery:
                self.view.display_error("Perfumery not found")
                return self.view.redirect_to_list()

            out_of_stock_items = self.inventory_repository.get_out_of_stock_by_perfumery(perfumery_id)

            if not out_of_stock_items:
                self.view.display_error("No out of stock items to export")
                return self.view.redirect_to_detail(perfumery_id)

            perfumery_name = perfumery.name.replace(" ", "_").lower()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            if format_type == 'csv':
                file_path = os.path.join(self.export_dir, f"out_of_stock_{perfumery_name}_{timestamp}.csv")
                self._export_to_csv(out_of_stock_items, file_path)
            elif format_type == 'doc':
                file_path = os.path.join(self.export_dir, f"out_of_stock_{perfumery_name}_{timestamp}.docx")
                self._export_to_doc(out_of_stock_items, perfumery, file_path)
            else:
                self.view.display_error("Invalid export format. Must be 'csv' or 'doc'.")
                return self.view.redirect_to_detail(perfumery_id)

            return self.view.send_export_file(file_path)

        except Exception as e:
            self.view.display_error(f"Error exporting data: {str(e)}")
            return self.view.redirect_to_detail(perfumery_id)

    def _export_to_csv(self, items, file_path):
        with open(file_path, 'w', newline='') as csvfile:
            fieldnames = ['Perfume ID', 'Perfume Name', 'Brand', 'Price']
            writer = csv.DictWriter(csvfile, fieldnames=fieldnames)

            writer.writeheader()
            for item in items:
                writer.writerow({
                    'Perfume ID': item.perfume.id,
                    'Perfume Name': item.perfume.name,
                    'Brand': item.perfume.brand,
                    'Price': item.perfume.price
                })

    def _export_to_doc(self, items, perfumery, file_path):
        doc = Document()
        doc.add_heading(f'Out of Stock Perfumes at {perfumery.name}', 0)
        doc.add_paragraph(f'Report generated on {datetime.now().strftime("%Y-%m-%d %H:%M")}')

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Perfume ID'
        header_cells[1].text = 'Perfume Name'
        header_cells[2].text = 'Brand'
        header_cells[3].text = 'Price'

        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.perfume.id)
            row_cells[1].text = item.perfume.name
            row_cells[2].text = item.perfume.brand
            row_cells[3].text = f"${item.perfume.price:.2f}"

        doc.add_paragraph('')
        doc.add_paragraph('Please restock these items as soon as possible.')

        doc.save(file_path)

    def _convert_perfumery_to_dict(self, perfumery) -> Dict[str, Any]:
        return {
            'id': perfumery.id,
            'name': perfumery.name,
            'address': perfumery.address,
            'city': perfumery.city,
            'phone': perfumery.phone,
            'email': perfumery.email,
            'manager_name': perfumery.manager_name
        }

    def _convert_inventory_to_dict(self, inventory) -> Dict[str, Any]:
        image_url = None
        if inventory.perfume.image_path:
            image_url = f"/static/images/perfumes/{os.path.basename(inventory.perfume.image_path)}"

        return {
            'id': inventory.id,
            'perfume_id': inventory.perfume.id,
            'perfume_name': inventory.perfume.name,
            'perfume_brand': inventory.perfume.brand,
            'perfume_price': inventory.perfume.price,
            'quantity': inventory.quantity,
            'is_available': inventory.is_available,
            'image_path': image_url
        }