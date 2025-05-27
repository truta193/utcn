import csv
from tkinter import messagebox, ttk, filedialog

from PIL._tkinter_finder import tk
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import numpy as np


class PerfumeryController:

    def __init__(self, perfumery_repository, inventory_repository, perfume_repository, view):

        self.perfumery_repository = perfumery_repository
        self.inventory_repository = inventory_repository
        self.perfume_repository = perfume_repository
        self.view = view

        self._gender_fig = None
        self._gender_canvas = None
        self._brand_fig = None
        self._brand_canvas = None

        self.perfumery_repository.add_observer(self)

        self.view._lang_service.add_listener(self)

        self._setup_view()

        self.refresh_view()

    def _setup_view(self):

        self.view.add_perfumery_button.config(command=self.handle_add_perfumery)
        self.view.update_perfumery_button.config(command=self.handle_update_perfumery)
        self.view.delete_perfumery_button.config(command=self.handle_delete_perfumery)
        self.view.clear_perfumery_button.config(command=self.handle_clear_perfumery)
        self.view.export_csv_button.config(command=self.handle_export_csv)
        self.view.export_doc_button.config(command=self.handle_export_doc)

        #        self.view.perfumery_search_button.config(command=self.handle_search_perfumery)

        self.view.perfumeries_table.bind("<<TreeviewSelect>>", lambda e: self.handle_perfumery_selection())
        self.view.inventory_summary_table.bind("<<TreeviewSelect>>",
                                               lambda e: self.handle_inventory_summary_selection())

        self.view.perfumery_form_canvas.bind_all("<MouseWheel>", self.handle_perfumery_form_mousewheel)

        self.update_translations()

    def refresh_view(self):

        perfumeries = self.perfumery_repository.get_all()
        self.update_perfumeries_list(perfumeries)

    def update_perfumeries_list(self, perfumeries):

        self.view._perfumeries = perfumeries

        for item in self.view.perfumeries_table.get_children():
            self.view.perfumeries_table.delete(item)

        if not perfumeries:
            self.view.perfumeries_table.insert('', 'end', values=('',
                                                                  self.view._lang_service.get_text("no_perfumeries"),
                                                                  '', ''))
            return

        for perfumery in perfumeries:
            values = (
                perfumery.id,
                perfumery.name,
                perfumery.city,
                perfumery.phone or ""
            )
            self.view.perfumeries_table.insert('', 'end', values=values)

    #

    #

    #

    def update_inventory_summary(self, perfumery_id):

        for item in self.view.inventory_summary_table.get_children():
            self.view.inventory_summary_table.delete(item)

        inventory_items = self.inventory_repository.get_all_by_perfumery(perfumery_id)

        if not inventory_items:
            self.view.inventory_summary_table.insert('', 'end', values=('',
                                                                        self.view._lang_service.get_text(
                                                                            "no_inventory_items"),
                                                                        '', '', ''))
            return

        for item in inventory_items:

            perfume = self.perfume_repository.get_by_id(item.perfume_id)

            if perfume:
                availability_text = self.view._lang_service.get_text(
                    "available") if item.is_available else self.view._lang_service.get_text("out_of_stock")

                values = (
                    perfume.id,
                    perfume.name,
                    perfume.brand,
                    item.quantity,
                    availability_text
                )
                self.view.inventory_summary_table.insert('', 'end', values=values)

    def handle_perfumery_selection(self):

        selected_items = self.view.perfumeries_table.selection()
        if not selected_items:
            return

        item_values = self.view.perfumeries_table.item(selected_items[0], 'values')
        perfumery_id = int(item_values[0])

        selected_perfumery = None
        for perfumery in self.view._perfumeries:
            if perfumery.id == perfumery_id:
                selected_perfumery = perfumery
                break

        if selected_perfumery:
            self.view._current_perfumery_id = selected_perfumery.id
            self.view.perfumery_name_var.set(selected_perfumery.name)
            self.view.address_var.set(selected_perfumery.address)
            self.view.city_var.set(selected_perfumery.city)
            self.view.phone_var.set(selected_perfumery.phone or "")
            self.view.email_var.set(selected_perfumery.email or "")
            self.view.manager_var.set(selected_perfumery.manager_name or "")

            self.update_inventory_summary(selected_perfumery.id)
            self.update_statistics(selected_perfumery.id)
            self.view.perfumery_form_canvas.yview_moveto(0)

    def handle_inventory_summary_selection(self):

        pass

    def handle_add_perfumery(self):

        name = self.view.perfumery_name_var.get().strip()
        address = self.view.address_var.get().strip()
        city = self.view.city_var.get().strip()
        phone = self.view.phone_var.get().strip() or None
        email = self.view.email_var.get().strip() or None
        manager_name = self.view.manager_var.get().strip() or None

        if not name or not address or not city:
            self.show_error("fill_required_perfumery")
            return

        self.perfumery_repository.create(
            name=name,
            address=address,
            city=city,
            phone=phone,
            email=email,
            manager_name=manager_name
        )

        self.handle_clear_perfumery()

    def handle_update_perfumery(self):

        if not self.view._current_perfumery_id:
            return

        name = self.view.perfumery_name_var.get().strip()
        address = self.view.address_var.get().strip()
        city = self.view.city_var.get().strip()
        phone = self.view.phone_var.get().strip() or None
        email = self.view.email_var.get().strip() or None
        manager_name = self.view.manager_var.get().strip() or None

        if not name or not address or not city:
            self.show_error("fill_required_perfumery")
            return

        self.perfumery_repository.update(
            self.view._current_perfumery_id,
            name=name,
            address=address,
            city=city,
            phone=phone,
            email=email,
            manager_name=manager_name
        )

    def handle_delete_perfumery(self):

        if not self.view._current_perfumery_id:
            return

        if not self.confirm_delete("confirm_delete_perfumery", "confirm_delete_perfumery_msg"):
            return

        inventory_items = self.inventory_repository.get_all_by_perfumery(self.view._current_perfumery_id)
        for item in inventory_items:
            self.inventory_repository.delete(item.id)

        self.perfumery_repository.delete(self.view._current_perfumery_id)

        self.handle_clear_perfumery()

    def handle_clear_perfumery(self):

        self.view._current_perfumery_id = None
        self.view.perfumery_name_var.set("")
        self.view.address_var.set("")
        self.view.city_var.set("")
        self.view.phone_var.set("")
        self.view.email_var.set("")
        self.view.manager_var.set("")

        for item in self.view.inventory_summary_table.get_children():
            self.view.inventory_summary_table.delete(item)

        for item in self.view.perfumeries_table.selection():
            self.view.perfumeries_table.selection_remove(item)

    def handle_search_perfumery(self):

        search_term = self.view.perfumery_search_var.get().strip()

        if not search_term:

            perfumeries = self.perfumery_repository.get_all()
        else:

            perfumeries = self.perfumery_repository.search_by_name(search_term)

        self.update_perfumeries_list(perfumeries)

    def handle_filter_city(self):

        city = self.view.filter_city_var.get()
        all_cities_text = self.view._lang_service.get_text("all_cities")

        perfumeries = self.perfumery_repository.get_all()

        self.update_perfumeries_list(perfumeries)

    def handle_perfumery_form_mousewheel(self, event):

        self.view.perfumery_form_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def update(self, observable, *args, **kwargs):

        if observable != self.perfumery_repository:
            return

        if len(args) > 0:
            event_type = args[0]

            if event_type in ["create", "update", "delete"]:
                perfumeries = self.perfumery_repository.get_all()
                self.update_perfumeries_list(perfumeries)

                cities = sorted(set(p.city for p in perfumeries))
                #                self.update_cities_filter(cities)

                message_key = f"{event_type}_perfumery_success"
                self.show_message(message_key)

    def on_language_changed(self, language_code):

        self.update_translations()

    def update_translations(self):

        self.view.perfumery_name_label.configure(text=self.view._lang_service.get_text("perfumery_name"))
        self.view.address_label.configure(text=self.view._lang_service.get_text("address"))
        self.view.city_label.configure(text=self.view._lang_service.get_text("city"))
        self.view.phone_label.configure(text=self.view._lang_service.get_text("phone"))
        self.view.email_label.configure(text=self.view._lang_service.get_text("email"))
        self.view.manager_label.configure(text=self.view._lang_service.get_text("manager"))
        self.view.inventory_summary_label.configure(text=self.view._lang_service.get_text("inventory_summary"))

        self.view.add_perfumery_button.configure(text=self.view._lang_service.get_text("add"))
        self.view.update_perfumery_button.configure(text=self.view._lang_service.get_text("update"))
        self.view.delete_perfumery_button.configure(text=self.view._lang_service.get_text("delete"))
        self.view.clear_perfumery_button.configure(text=self.view._lang_service.get_text("clear"))
        self.view.export_csv_button.configure(text=self.view._lang_service.get_text("export_csv"))
        self.view.export_doc_button.configure(text=self.view._lang_service.get_text("export_doc"))

        #        self.view.perfumery_search_label.configure(text=self.view._lang_service.get_text("search"))
        # self.view.perfumery_search_button.configure(text=self.view._lang_service.get_text("search"))

        self.view.perfumeries_table.heading('name', text=self.view._lang_service.get_text("name"))
        self.view.perfumeries_table.heading('city', text=self.view._lang_service.get_text("city"))
        self.view.perfumeries_table.heading('phone', text=self.view._lang_service.get_text("phone"))

        self.view.inventory_summary_table.heading('name', text=self.view._lang_service.get_text("name"))
        self.view.inventory_summary_table.heading('brand', text=self.view._lang_service.get_text("brand"))
        self.view.inventory_summary_table.heading('quantity', text=self.view._lang_service.get_text("quantity"))
        self.view.inventory_summary_table.heading('available', text=self.view._lang_service.get_text("available"))

        self.view.stats_label.configure(text=self.view._lang_service.get_text("statistics"))
        self.view.gender_chart_frame.configure(text=self.view._lang_service.get_text("gender_distribution"))
        self.view.brand_chart_frame.configure(text=self.view._lang_service.get_text("brand_distribution"))

        try:

            for window_id in self.view.perfumery_form_canvas.find_all():
                if self.view.perfumery_form_canvas.type(window_id) == "window":
                    window_info = self.view.perfumery_form_canvas.itemconfig(window_id)
                    window = window_info.get("window", None)[4]
                    if isinstance(window, ttk.LabelFrame):
                        window.configure(text=self.view._lang_service.get_text("perfumery_details"))
                        break
        except Exception as e:
            print(f"Error updating perfumery form label: {e}")

        if self.view._current_perfumery_id:
            self.update_inventory_summary(self.view._current_perfumery_id)

    def show_message(self, message_key):

        message = self.view._lang_service.get_text(message_key)
        self.view.show_message(message)

    def show_error(self, message_key):

        message = self.view._lang_service.get_text(message_key)
        messagebox.showerror(self.view._lang_service.get_text("error"), message)

    def confirm_delete(self, title_key, message_key):

        return messagebox.askyesno(
            self.view._lang_service.get_text(title_key),
            self.view._lang_service.get_text(message_key)
        )

    def update_statistics(self, perfumery_id):

        plt.close('all')

        for widget in self.view.gender_chart_frame.winfo_children():
            widget.destroy()
        for widget in self.view.brand_chart_frame.winfo_children():
            widget.destroy()

        inventory_items = self.inventory_repository.get_all_by_perfumery(perfumery_id)

        if not inventory_items:
            no_data_label1 = ttk.Label(self.view.gender_chart_frame,
                                       text=self.view._lang_service.get_text("no_data"),
                                       font=('Arial', 12))
            no_data_label1.pack(pady=20, expand=True)

            no_data_label2 = ttk.Label(self.view.brand_chart_frame,
                                       text=self.view._lang_service.get_text("no_data"),
                                       font=('Arial', 12))
            no_data_label2.pack(pady=20, expand=True)
            return

        perfumes = []
        for item in inventory_items:
            perfume = self.perfume_repository.get_by_id(item.perfume_id)
            if perfume:
                perfumes.append(perfume)

        if not perfumes:
            no_data_label1 = ttk.Label(self.view.gender_chart_frame,
                                       text=self.view._lang_service.get_text("no_data"),
                                       font=('Arial', 12))
            no_data_label1.pack(pady=20, fill=tk.BOTH, expand=True)

            no_data_label2 = ttk.Label(self.view.brand_chart_frame,
                                       text=self.view._lang_service.get_text("no_data"),
                                       font=('Arial', 12))
            no_data_label2.pack(pady=20, fill=tk.BOTH, expand=True)
            return

        self.view.gender_chart_frame.pack(expand=True, pady=5, ipady=100)
        self.view.brand_chart_frame.pack(expand=True, pady=5, ipady=100)

        self.create_gender_chart(perfumes)
        self.create_brand_chart(perfumes)

    def create_gender_chart(self, perfumes):

        gender_counts = {"male": 0, "female": 0, "unisex": 0, "unknown": 0}

        for perfume in perfumes:
            if not perfume.gender:
                gender_counts["unknown"] += 1
            elif perfume.gender.lower() in gender_counts:
                gender_counts[perfume.gender.lower()] += 1
            else:
                gender_counts["unknown"] += 1

        gender_data = {k: v for k, v in gender_counts.items() if v > 0}

        if not gender_data:
            ttk.Label(self.view.gender_chart_frame,
                      text=self.view._lang_service.get_text("no_data"),
                      font=('Arial', 12)).pack(pady=20)
            return

        fig = plt.figure(figsize=(6, 4), dpi=100, facecolor='white')
        ax = fig.add_subplot(111)

        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#c2c2f0']

        wedges, texts, autotexts = ax.pie(
            gender_data.values(),
            labels=[self.view._lang_service.get_text(k) for k in gender_data.keys()],
            autopct='%1.1f%%',
            startangle=90,
            colors=colors,
            textprops={'color': 'black', 'fontsize': 10, 'weight': 'bold'}
        )

        ax.set_facecolor('white')

        ax.set_title(self.view._lang_service.get_text("gender_distribution"),
                     color='black', fontsize=12, weight='bold')

        chart_container = ttk.Frame(self.view.gender_chart_frame)
        chart_container.pack(expand=True, padx=10, pady=10)

        from matplotlib.backends.backend_tkagg import NavigationToolbar2Tk

        canvas = FigureCanvasTkAgg(fig, master=chart_container)
        canvas.draw()

        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack(expand=True)

        toolbar = NavigationToolbar2Tk(canvas, chart_container)
        toolbar.update()

        self._gender_fig = fig
        self._gender_canvas = canvas

    def create_brand_chart(self, perfumes):

        brand_counts = {}
        for perfume in perfumes:
            brand = perfume.brand if perfume.brand else "Unknown"
            if brand in brand_counts:
                brand_counts[brand] += 1
            else:
                brand_counts[brand] = 1

        if not brand_counts:
            ttk.Label(self.view.brand_chart_frame,
                      text=self.view._lang_service.get_text("no_data"),
                      font=('Arial', 12)).pack(pady=20)
            return

        if len(brand_counts) > 5:
            sorted_brands = sorted(brand_counts.items(), key=lambda x: x[1], reverse=True)
            top_brands = dict(sorted_brands[:4])
            others_count = sum(dict(sorted_brands[4:]).values())
            top_brands[self.view._lang_service.get_text("others")] = others_count
            brand_counts = top_brands

        fig = plt.figure(figsize=(6, 4), dpi=100, facecolor='white')
        ax = fig.add_subplot(111)

        colors = ['#ff9999', '#66b3ff', '#99ff99', '#ffcc99', '#c2c2f0', '#ff6666', '#99ccff']

        wedges, texts, autotexts = ax.pie(
            brand_counts.values(),
            labels=brand_counts.keys(),
            autopct='%1.1f%%',
            startangle=90,
            colors=colors,
            textprops={'color': 'black', 'fontsize': 10, 'weight': 'bold'}
        )

        ax.set_facecolor('white')

        ax.set_title(self.view._lang_service.get_text("brand_distribution"),
                     color='black', fontsize=12, weight='bold')

        chart_container = ttk.Frame(self.view.brand_chart_frame)
        chart_container.pack(expand=True, padx=10, pady=10)

        from matplotlib.backends.backend_tkagg import NavigationToolbar2Tk

        canvas = FigureCanvasTkAgg(fig, master=chart_container)
        canvas.draw()

        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack(expand=True)

        toolbar = NavigationToolbar2Tk(canvas, chart_container)
        toolbar.update()

        self._brand_fig = fig
        self._brand_canvas = canvas

    def handle_export_doc(self):

        if not self.view._current_perfumery_id:
            self.show_error("select_perfumery_first")
            return

        selected_perfumery = None
        for perfumery in self.view._perfumeries:
            if perfumery.id == self.view._current_perfumery_id:
                selected_perfumery = perfumery
                break

        if not selected_perfumery:
            return

        inventory_items = self.inventory_repository.get_out_of_stock_by_perfumery(self.view._current_perfumery_id)

        if not inventory_items:
            self.show_message("no_out_of_stock")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialfile=f"{selected_perfumery.name}_out_of_stock.docx"
        )

        if not file_path:
            return

        try:

            doc = Document()

            doc.add_heading(f"{selected_perfumery.name}: {self.view._lang_service.get_text('out_of_stock_items')}", 0)

            doc.add_paragraph(
                f"{self.view._lang_service.get_text('out_of_stock_report')} {selected_perfumery.name}, {selected_perfumery.address}, {selected_perfumery.city}")

            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            header_cells = table.rows[0].cells
            header_cells[0].text = "ID"
            header_cells[1].text = self.view._lang_service.get_text("perfume")
            header_cells[2].text = self.view._lang_service.get_text("brand")
            header_cells[3].text = self.view._lang_service.get_text("quantity")

            for item in inventory_items:

                perfume = self.perfume_repository.get_by_id(item.perfume_id)
                if perfume:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(perfume.id)
                    row_cells[1].text = perfume.name
                    row_cells[2].text = perfume.brand
                    row_cells[3].text = str(item.quantity)

            doc.save(file_path)

            self.show_message("export_doc_success")

        except Exception as e:
            print(f"Error exporting to DOC: {e}")
            self.show_error("export_error")

    def handle_export_csv(self):

        if not self.view._current_perfumery_id:
            self.show_error("select_perfumery_first")
            return

        selected_perfumery = None
        for perfumery in self.view._perfumeries:
            if perfumery.id == self.view._current_perfumery_id:
                selected_perfumery = perfumery
                break

        if not selected_perfumery:
            return

        inventory_items = self.inventory_repository.get_out_of_stock_by_perfumery(self.view._current_perfumery_id)

        if not inventory_items:
            self.show_message("no_out_of_stock")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")],
            initialfile=f"{selected_perfumery.name}_out_of_stock.csv"
        )

        if not file_path:
            return

        try:

            with open(file_path, 'w', newline='') as csvfile:
                writer = csv.writer(csvfile)

                writer.writerow([
                    "ID",
                    self.view._lang_service.get_text("perfume"),
                    self.view._lang_service.get_text("brand"),
                    self.view._lang_service.get_text("quantity")
                ])

                for item in inventory_items:

                    perfume = self.perfume_repository.get_by_id(item.perfume_id)
                    if perfume:
                        writer.writerow([
                            perfume.id,
                            perfume.name,
                            perfume.brand,
                            item.quantity
                        ])

            self.show_message("export_csv_success")

        except Exception as e:
            print(f"Error exporting to CSV: {e}")
            self.show_error("export_error")
